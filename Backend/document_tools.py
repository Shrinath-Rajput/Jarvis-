# document_tools.py
"""
Document and document generation tools for JARVIS
"""

import os
import json
from datetime import datetime


class DocumentTools:
    """Handle document operations"""
    
    @staticmethod
    def create_resume(name, email, phone, experience, education, skills, save_path=None):
        """Create a resume document"""
        try:
            if save_path is None:
                save_path = os.path.expanduser("~/Documents/resume.docx")
            
            save_path = os.path.expanduser(save_path)
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading(name, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact info
            contact = doc.add_paragraph(f"{email} | {phone}")
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Experience
            if experience:
                doc.add_heading("Experience", level=2)
                for exp in experience:
                    doc.add_paragraph(f"{exp['title']} at {exp['company']}", style='List Bullet')
                    doc.add_paragraph(f"{exp['description']}", style='List Bullet 2')
            
            # Education
            if education:
                doc.add_heading("Education", level=2)
                for edu in education:
                    doc.add_paragraph(f"{edu['degree']} from {edu['school']}", style='List Bullet')
            
            # Skills
            if skills:
                doc.add_heading("Skills", level=2)
                skills_text = ", ".join(skills)
                doc.add_paragraph(skills_text)
            
            doc.save(save_path)
            
            return {"success": True, "message": f"Resume created at {save_path}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def create_cover_letter(name, company, position, body, save_path=None):
        """Create a cover letter"""
        try:
            if save_path is None:
                save_path = os.path.expanduser("~/Documents/cover_letter.docx")
            
            save_path = os.path.expanduser(save_path)
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            
            # Date
            doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
            
            # Company address placeholder
            doc.add_paragraph("Company Address")
            doc.add_paragraph()
            
            # Greeting
            doc.add_paragraph(f"Dear Hiring Manager,")
            doc.add_paragraph()
            
            # Body
            doc.add_paragraph(f"I am writing to express my interest in the {position} position at {company}.")
            doc.add_paragraph(body)
            
            # Closing
            doc.add_paragraph()
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph(name)
            
            doc.save(save_path)
            
            return {"success": True, "message": f"Cover letter created at {save_path}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def spell_check(text):
        """Spell check text"""
        try:
            from pyspellchecker import SpellChecker
            
            spell = SpellChecker()
            misspelled = spell.unknown(text.split())
            
            corrections = {}
            for word in misspelled:
                corrections[word] = spell.correction(word)
            
            return {"success": True, "misspelled": list(misspelled), "corrections": corrections}
        except Exception as e:
            # Fallback if pyspellchecker not installed
            return {"success": False, "error": f"Spell check unavailable: {str(e)}"}
    
    @staticmethod
    def summarize_pdf(pdf_path):
        """Summarize a PDF document"""
        try:
            pdf_path = os.path.expanduser(pdf_path)
            
            if not os.path.exists(pdf_path):
                return {"success": False, "error": f"PDF not found: {pdf_path}"}
            
            from PyPDF2 import PdfReader
            import re
            
            reader = PdfReader(pdf_path)
            text = ""
            
            for page in reader.pages:
                text += page.extract_text()
            
            # Simple summarization - extract key sentences
            sentences = re.split(r'[.!?]', text)
            summary = ". ".join(sentences[:5]) + "."
            
            return {"success": True, "summary": summary}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def translate_text(text, target_language="spanish"):
        """Translate text"""
        try:
            # Using free translation API
            import requests
            
            language_codes = {
                "spanish": "es",
                "french": "fr",
                "german": "de",
                "italian": "it",
                "portuguese": "pt",
                "russian": "ru",
                "japanese": "ja",
                "chinese": "zh",
                "korean": "ko"
            }
            
            target_code = language_codes.get(target_language.lower(), "es")
            
            # Using simple approach with requests
            url = f"https://api.mymemory.translated.net/get?q={text}&langpair=en|{target_code}"
            response = requests.get(url)
            data = response.json()
            
            if data['responseStatus'] == 200:
                translated = data['responseData']['translatedText']
                return {"success": True, "translated": translated}
            else:
                return {"success": False, "error": "Translation failed"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def generate_report(title, sections, save_path=None):
        """Generate a report document"""
        try:
            if save_path is None:
                save_path = os.path.expanduser("~/Documents/report.docx")
            
            save_path = os.path.expanduser(save_path)
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            
            # Title
            doc.add_heading(title, level=1)
            
            # Date
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()
            
            # Sections
            for section in sections:
                doc.add_heading(section.get('title', 'Section'), level=2)
                doc.add_paragraph(section.get('content', ''))
            
            doc.save(save_path)
            
            return {"success": True, "message": f"Report created at {save_path}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def read_pdf(pdf_path):
        """Read and extract text from PDF"""
        try:
            pdf_path = os.path.expanduser(pdf_path)
            
            if not os.path.exists(pdf_path):
                return {"success": False, "error": f"PDF not found: {pdf_path}"}
            
            from PyPDF2 import PdfReader
            
            reader = PdfReader(pdf_path)
            text = ""
            
            for page in reader.pages:
                text += page.extract_text()
            
            return {"success": True, "text": text}
        except Exception as e:
            return {"success": False, "error": str(e)}


# Export
document_tools = DocumentTools()
